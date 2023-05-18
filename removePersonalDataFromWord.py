# -*- coding: utf-8 -*-
import os
from docx import Document
import spacy
import re
import pandas as pd


new_name = "new_file_name"
directory_path = "summarization-resume-vacancy-matching/docx_after_preprocessing/docx"


# Load the 'en_core_web_sm' model for NLP
nlp = spacy.load('en_core_web_sm')

def remove_personal_data(directory_path):
    for nameIndex, filename in enumerate(os.listdir(directory_path)):
        
        if filename.endswith(".docx"):
           
            # Load the Word document
            doc = Document(os.path.join(directory_path, filename))
            
            paragraph = doc.paragraphs[0]
            
          

            # Define regular expressions for phone numbers, email addresses, URLs, and addresses
            phone_regex = re.compile(r"(\+)?([0-9]{1,3})?( )?(\([0-9]{1,3}\))?( )?[(\d+((\-\d+)+)]{10,15}")
            email_regex = re.compile(r'^[a-zA-Z]+(?:(?:_[a-zA-Z0-9]+)+\.[A-Za-z0-9]+|\.[a-zA-Z][a-zA-Z0-9]*)?@(?:[a-zA-Z0-9]+\.)*[a-zA-Z0-9]{2,}$')
            url_regex = re.compile(r'http://\S+|https://\S+')
            address_regex = re.compile(r'\d+ [\w\s]+, [\w\s]+, [\w\s]+')
            fullName_regex = re.compile(paragraph.text)
           
            url_regex = re.compile(r'http://\S+|https://\S+')

            # Loop through each paragraph in the document
            for para in doc.paragraphs:
                # Parse the paragraph using the 'en_core_web_sm' model

                parsed_para = nlp(para.text)
                # print(parsed_para)
                # Loop through each entity in the parsed paragraph
                for ent in parsed_para.ents:
                    # Check if the entity is a phone number, email address, URL, or address
                    if ent.label_ == 'PHONE':
                        # Replace the phone number with asterisks in the parsed paragraph
                        parsed_para = parsed_para[:ent.start_char] + \
                            '*' * len(ent.text) + parsed_para[ent.end_char:]
                    elif ent.label_ == 'Email:':
                        # Replace the email address with asterisks in the parsed paragraph
                        parsed_para = parsed_para[:ent.start_char] + \
                            '*' * len(ent.text) + parsed_para[ent.end_char:]
                    elif ent.label_ == 'URL':
                        # Replace the URL with asterisks in the parsed paragraph
                        parsed_para = parsed_para[:ent.start_char] + \
                            '*' * len(ent.text) + parsed_para[ent.end_char:]
                    elif ent.label_ == 'ADDRESS':
                        # Replace the address with asterisks in the parsed paragraph
                        parsed_para = parsed_para[:ent.start_char] + \
                            '*' * len(ent.text) + parsed_para[ent.end_char:]

                # Update the text of the original paragraph with the modified text from the parsed paragraph
                para.text = parsed_para.text
                if fullName_regex.search(para.text):
                    para.text = fullName_regex.sub(
                        '*' * len(fullName_regex.search(para.text).group()), para.text)
                # Check if the paragraph contains a phone number, email address, URL, or address using regular expressions
                if phone_regex.search(para.text):
                    # Replace the phone number with asterisks
                    para.text = phone_regex.sub('*' * 10, para.text)
                if email_regex.search(para.text):

                    # Replace the email address with asterisks
                    para.text = email_regex.sub(
                        '*' * len(email_regex.search(para.text).group()), para.text)
                if url_regex.search(para.text):
                    # Replace the URL with asterisks
                    para.text = url_regex.sub(
                        '*' * len(url_regex.search(para.text).group()), para.text)
                if address_regex.search(para.text):
                    # Replace the address with asterisks
                    para.text = address_regex.sub(
                        '*' * len(address_regex.search(para.text).group()), para.text)
            # Save the modified document
            #doc.save('modified_cv.docx')


            email_regex1 =r'(email|Email|E-mail)(:)*( )*[A-Za-z0-9]*.*@[A-Za-z]*\.?[A-Za-z0-9]*'
            email_regex2 = r'[A-Za-z0-9]*.*@[A-Za-z]*\.?[A-Za-z0-9]*'
            date_regex1 =r"\d{2}.\d{2}.\d{4}"
            birth_redex1 =r"Birthday(:)*( )*(\d{4}|(\d{2}.\d{2}.\d{4}))|Date of birth(:)*( )*(\d{2}.\d{2}.\d{4})"
            #year_redex1 = r"\d{4}"
            phone_regex1 = r"(phone|Tel|Phone|Mobile phone)(:)*( )*(((\+)?([0-9]{1,3})?( )?(\([0-9]{1,3}\))?( )?[(\d+((\-\d+)+)]{10,15})|(\d{3} \d{3} \d{2} \d{2}))"
            phone_regex2 = r"((\+)?([0-9]{1,3})?( )?(\([0-9]{1,3}\))?( )?[(\d+((\-\d+)+)]{10,15})|(\d{3} \d{3} \d{2} \d{2})"
            url_regex1 = r'(Linkedln|linkedIn|Github)(:)*( )*http(s)*://\S+|https://\S+' 
            url_regex2 = r'http://\S+|https://\S+'
            age_regex1 = r'Age:( )*\d{2} y.o.'
            #city_regex1 = r'(City of Residence|Residence|location)(:)*( )*'
            #citis_df = pd.read_csv('citis.csv')
            #citis_names = citis_df['city'].values.tolist()
            #citis_names.extend(['Rishon LeZion', 'Herzliya','Tel Aviv', 'Tel-Aviv','Beer Sheva','Hod-Hasharon', 'Israel', 'Area'])

            # doc = docx.Document('./CV Dmitry Itskov java_FullStackDeveloper.docx')
            # doc = docx.Document('Rezunenko Timur. Web-developer.docx');
            # doc = docx.Document('./CV_Vladislav_Feldsher_frontend_developer.docx');
            #doc = docx.Document('summarization-resume-vacancy-matching/docx_after_preprocessing/docx/CV Dmitry Itskov java_FullStackDeveloper.docx');
            for table in doc.tables:
            # Loop through each row in the table
                for row in table.rows:
                # Loop through each cell in the row
                    for cell in row.cells:
                        cell.text = re.sub(phone_regex1, '', cell.text).replace('\n', ' ') 
                        cell.text = re.sub(phone_regex2, '', cell.text).replace('\n', ' ')
                        cell.text = re.sub(email_regex1, '', cell.text).replace('\n', ' ')
                        cell.text = re.sub(email_regex2, '', cell.text).replace('\n', ' ') 
                        cell.text = re.sub(url_regex1, '', cell.text).replace('\n', ' ')
                        cell.text = re.sub(url_regex2, '', cell.text).replace('\n', ' ')
                        cell.text = re.sub(age_regex1, '', cell.text).replace('\n', ' ')
                        cell.text = re.sub(birth_redex1, '', cell.text).replace('\n', ' ')
                        cell.text = re.sub(date_regex1, '', cell.text).replace('\n', ' ') 
                        #cell.text = re.sub(year_redex1, '', cell.text).replace('\n', ' ')
                        #for i in citis_names:
                            #cell.text = re.sub(city_regex1+i, '', cell.text).replace('\n', ' ');
                            #cell.text = re.sub(i, '', cell.text).replace('\n', ' ');
                    #print(cell.text)    
            

            for para in doc.paragraphs:
                parsed_para = nlp(para.text)
                para.text = parsed_para.text
                para.text = re.sub(phone_regex1, '', para.text).replace('\n', ' ') 
                para.text = re.sub(phone_regex2, '', para.text).replace('\n', ' ')
                para.text = re.sub(email_regex1, '', para.text).replace('\n', ' ')
                para.text = re.sub(url_regex1, '', para.text).replace('\n', ' ')
                para.text = re.sub(age_regex1, '', para.text).replace('\n', ' ')
                para.text = re.sub(birth_redex1, '', para.text).replace('\n', ' ')
                para.text = re.sub(date_regex1, '', para.text).replace('\n', ' ') 
                #para.text = re.sub(year_redex1, '', para.text).replace('\n', ' ')
                #for i in citis_names:
                    #para.text = re.sub(city_regex1+i, '', para.text).replace('\n', ' ');
                    #para.text = re.sub(i, '', para.text).replace('\n', ' ');
            
            # Save the modified document
            
            new_filename = f"{nameIndex+1}.docx"
            doc.save(os.path.join('results', new_filename))
# Call the function with the file path of the CV
#remove_personal_data('summarization-resume-vacancy-matching/docx_after_preprocessing/docx/CV Dmitry Itskov java_FullStackDeveloper.docx')
remove_personal_data(directory_path)



