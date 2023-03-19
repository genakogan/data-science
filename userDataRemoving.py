from docx import Document
import spacy
import re
import os
# Load the 'en_core_web_sm' model for NLP
nlp = spacy.load('en_core_web_sm')


directory_path = "summarization-resume-vacancy-matching/docx_after_preprocessing/docx"

# Set the new name you want to use for the files
new_name = "new_file_name"
def remove_personal_data(doc):
    for i, filename in enumerate(os.listdir(directory_path)):

        if filename.endswith(".docx"):
            # Open the file with docx module
            doc = Document(os.path.join(directory_path, filename))
            
            # Load the Word document
            #doc = docx.Document(filepath)
            """  print(doc.paragraphs[1].text)
            doc.paragraphs[1].clear()  # Clear the first paragraph
            doc.save('test.docx') """
            
            
            first_word = None
            z = 0
            while z < len(doc.paragraphs) and not first_word:
                paragraph = doc.paragraphs[z]
                if len(paragraph.text.strip()) > 0:
                    words = paragraph.text.strip().split()
                    for word in words:
                        if len(word.strip()) > 0:
                            first_word = word
                            break
                z += 1
            # print first word in docx
            #print(paragraph.text)

            # Define regular expressions for phone numbers, email addresses, URLs, and addresses
            phone_regex = re.compile(r'(\+?\d{1,2}[- ]?)?\d{3}[- ]?\d{3}[- ]?\d{4}')
            email_regex = re.compile(r'[\w\.-]+@[\w\.-]+')
            url_regex = re.compile(r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+')
            address_regex = re.compile(r'\d+ [\w\s]+, [\w\s]+, [\w\s]+')
            authors_name_regex = re.compile(paragraph.text)
        
            tableData=[]
            for table in doc.tables:
            # Loop through each row in the table
                for row in table.rows:
                    # Loop through each cell in the row
                    for cell in row.cells:
                        # Parse the cell text using the 'en_core_web_sm' model
                        
                        parsed_cell = nlp(cell.text)
                        tableData.append(parsed_cell.text)
                        
        
            
            # Get the first table in the document
            """ table = doc.tables[0]
            
            # Get the first cell in the first row of the table
            cell = table.cell(0, 2)
            # Check if the cell text contains the text to remove
            if '055 940 94 52' in cell.text:
                # Remove the text from the cell text
                new_text = cell.text.replace('055 940 94 52', '****')

                # Replace the cell text with the modified text
                cell.text = new_text """
            

            
            # Loop through each paragraph in the document
            for para in doc.paragraphs:
                # Parse the paragraph using the 'en_core_web_sm' model
                
            
                parsed_para = nlp(para.text)
                """ for token in parsed_para:
                    print(token.text, token.lemma_, token.pos_, token.tag_, token.dep_,
                    token.shape_, token.is_alpha, token.is_stop, token.idx, token.head.text) """
                # Loop through each entity in the parsed paragraph
                for ent in parsed_para.ents:
                    #print(ent.text, ent.label_)
                    # Check if the entity is a phone number, email address, URL, or address
                    if ent.label_ == 'PHONE':
                        # Replace the phone number with asterisks in the parsed paragraph
                        parsed_para = parsed_para[:ent.start_char] + '*' * len(ent.text) + parsed_para[ent.end_char:]
                    elif ent.label_ == 'email':
                        
                        # Replace the email address with asterisks in the parsed paragraph
                        parsed_para = parsed_para[:ent.start_char] + '*' * len(ent.text) + parsed_para[ent.end_char:]
                    elif ent.label_ == 'URL':
                        # Replace the URL with asterisks in the parsed paragraph
                        parsed_para = parsed_para[:ent.start_char] + '*' * len(ent.text) + parsed_para[ent.end_char:]
                    elif ent.label_ == 'ADDRESS':
                        # Replace the address with asterisks in the parsed paragraph
                        parsed_para = parsed_para[:ent.start_char] + '*' * len(ent.text) + parsed_para[ent.end_char:]

            

                # Update the text of the original paragraph with the modified text from the parsed paragraph
                para.text = parsed_para.text

                # Check if the paragraph contains a phone number, email address, URL, or address using regular expressions
                if authors_name_regex.search(para.text):
                    para.text = authors_name_regex.sub('*' * len(authors_name_regex.search(para.text).group()), para.text)
                if phone_regex.search(para.text):
                    # Replace the phone number with asterisks
                    para.text = phone_regex.sub('*' * 10, para.text)
                if email_regex.search(para.text):
                    # Replace the email address with asterisks
                    para.text = email_regex.sub('*' * len(email_regex.search(para.text).group()), para.text)
                if url_regex.search(para.text):
                    # Replace the URL with asterisks
                    para.text = url_regex.sub('*' * len(url_regex.search(para.text).group()), para.text)
                if address_regex.search(para.text):
                    # Replace the address with asterisks
                    para.text = address_regex.sub('*' * len(address_regex.search(para.text).group()), para.text)

            # Save the modified document
            new_filename = f"{new_name}{i}.docx"
            doc.save(os.path.join('results', new_filename))
        #remove_personal_data(doc)
# Call the function with the file path of the CV
remove_personal_data(directory_path)


